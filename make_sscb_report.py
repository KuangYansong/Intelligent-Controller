from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_SCRIPT_DIR = Path(
    r"C:\Users\了女哦\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents\scripts"
)
if str(SKILL_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SKILL_SCRIPT_DIR))

from table_geometry import apply_table_geometry, exact_column_widths  # noqa: E402


OUT = Path(r"C:\Users\了女哦\Desktop\固态产品\设计报告\智能固态断路器控制器软件技术报告.docx")


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Microsoft YaHei", size=11, color=None, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style._element.rPr.rFonts.set(qn("w:cs"), name)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_paragraph_format(p, *, before=0, after=6, line=1.1, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text="", *, style=None, size=11, color=None, bold=False, italic=False, align=None, before=0, after=6, line=1.1, font="Microsoft YaHei"):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, before=before, after=after, line=line, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, name=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_formula(doc, text):
    return add_para(
        doc,
        text,
        size=11,
        font="Cambria Math",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=4,
        line=1.0,
    )


def set_cell_text(cell, text, *, bold=False, size=9.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.0, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths, *, header_fill="F2F4F7", font_size=9.5):
    widths_dxa = [int(round(w * 1440)) if isinstance(w, float) and w < 100 else int(w) for w in widths]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0].cells
    for idx, title in enumerate(headers):
        set_cell_text(header[idx], title, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        if header_fill:
            shade_cell(header[idx], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
    apply_table_geometry(
        table,
        exact_column_widths(widths_dxa, 9360),
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    return table


def add_section_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    set_paragraph_format(
        p,
        before=16 if level == 1 else 10 if level == 2 else 8,
        after=8 if level == 1 else 5 if level == 2 else 4,
        line=1.05,
    )
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=15.5, color=RGBColor(46, 116, 181), bold=True)
    elif level == 2:
        set_run_font(r, size=12.5, color=RGBColor(46, 116, 181), bold=True)
    else:
        set_run_font(r, size=11.5, color=RGBColor(31, 77, 120), bold=True)
    return p


def add_title_block(doc):
    add_para(
        doc,
        "智能固态断路器控制器软件技术报告",
        size=22,
        color=RGBColor(0, 0, 0),
        bold=True,
        before=0,
        after=3,
        line=1.0,
    )
    add_para(
        doc,
        "基于 TI TMS320F280049C 平台和当前工程源码的实现分析",
        size=11,
        color=RGBColor(85, 85, 85),
        italic=True,
        before=0,
        after=6,
        line=1.0,
    )
    meta = [
        ("报告依据", "当前项目源码、硬件网表、需求说明、通信协议、FRAM 地址映射和状态机说明"),
        ("项目范围", "E:\\DSP\\CCS\\ccs\\软件代码 与 E:\\DSP\\CCS\\ccs\\workspace\\F280049C\\bord"),
        ("编制日期", "2026-06-16"),
        ("适用对象", "项目管理、技术管理和联调人员"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        set_paragraph_format(p, before=0, after=1, line=1.0)
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, size=9.5, color=RGBColor(85, 85, 85), bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=9.5, color=RGBColor(85, 85, 85))
    add_para(
        doc,
        "本报告仅陈述当前工程中可以从代码和文档确认的实现；无法从项目文件确认的内容，统一按“未发现足够依据”处理。",
        size=9.5,
        color=RGBColor(96, 96, 96),
        italic=True,
        before=0,
        after=10,
        line=1.0,
    )


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 15.5, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 12.5, RGBColor(46, 116, 181), 10, 5),
        ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
    ]:
        st = styles[name]
        set_style_font(st, size=size, color=color, bold=True)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.05

    add_title_block(doc)

    add_section_heading(doc, "1. 项目概述", 1)
    add_para(
        doc,
        "当前工程面向智能固态断路器控制器的嵌入式软件实现，核心目标是把电压、电流和温度采样、短路和过载保护、CAN 通信、参数存储和故障记录整合为一条可运行的闭环链路。软件运行在 TI TMS320F280049C 上，借助 ADC、ePWM、CMPSS、CANB、SPIB 和 I2CA 这些外设完成实时控制与状态管理。",
    )
    add_para(
        doc,
        "从职责划分看，app 负责系统状态机和主调度，driver 负责外设初始化与读写抽象，protection 负责过流、过压和过温判断，protocol 负责 CAN 帧封装与命令处理，storage 负责参数镜像和故障记录结构，selftest 负责上电自检，common 则保存公共类型、参数和状态定义。",
    )
    add_table(
        doc,
        ["模块", "当前作用"],
        [
            ["app", "组织初始化、状态机、故障联动和主循环。"],
            ["driver", "封装 ADC、CAN、CMPSS、ePWM、I2C、SPI 和驱动器状态读取。"],
            ["protection", "完成 RMS、I2t、过压和过温判断。"],
            ["protocol", "完成实时帧、心跳帧、参数读写和手动复位命令。"],
            ["storage", "提供参数镜像、故障日志和 FRAM 访问接口。"],
            ["selftest", "进行轻量化上电自检和驱动器状态检查。"],
            ["common", "维护统一状态码、故障码、参数结构和时间基准。"],
        ],
        [1.25, 5.25],
    )
    add_para(
        doc,
        "软件当前没有引入 RTOS 或 DMA；应用层以 sscb_firmware_init 完成一次性初始化，以 sscb_firmware_run_once 作为主循环节拍入口。硬件上电后，系统先完成外设和参数初始化，再通过自检决定是否进入 READY 或 LOCKOUT，随后进入采样、换算、保护和 CAN 上报的持续运行流程。",
    )
    add_para(
        doc,
        "运行链路可概括为：上电初始化 -> 参数加载 -> 自检 -> ADC 采样 -> 工程量换算 -> 保护判断 -> 状态机切换 -> CAN 周期上报 -> 必要时写入故障记录。这个链路已经在当前工程中形成闭环，但 FRAM 持久化、时间同步和部分扩展报文仍以接口和数据结构形式预留。",
    )

    add_section_heading(doc, "2. 软件总体架构", 1)
    add_para(
        doc,
        "应用层的入口清晰分为两步：sscb_firmware_init 负责建立系统对象、时基、参数镜像、门极驱动配置和底层外设；sscb_firmware_run_once 负责读取最新 ADC 结果、换算工程量、执行保护、驱动状态机、发送周期报文并处理收到的 CAN 命令。当前可见代码中未发现独立的 ADC 中断服务函数，采样结果由 ePWM 触发后在主循环中读取最新结果寄存器。",
    )
    add_para(
        doc,
        "状态机实现于 app/system.c，覆盖 INIT、SELF_TEST、READY、RUN、WARN、FAULT_ACTIVE、RECOVER_WAIT、RECOVER_TRY 和 LOCKOUT。短路故障直接进入 LOCKOUT；过载故障在故障消失后可进入恢复等待，再进入恢复尝试；过温仅置为 WARN，并在温度降回恢复阈值后回到 RUN。上位机手动复位只对 LOCKOUT 有效，且必须同时满足 Trip 清除、驱动器就绪和故障条件解除。",
    )
    add_para(
        doc,
        "当前工程中，ADC、保护、CAN 和参数存储之间的数据关系比较直接：ADC 输出原始值，driver 层换算后得到工程量，protection 层基于工程量计算保护结果，system 层维护状态和故障日志，protocol 层读取状态字和参数结构生成报文，storage 层负责参数镜像与故障记录结构。"
    )
    add_formula(doc, "INIT -> SELF_TEST -> READY -> RUN")
    add_formula(doc, "RUN / WARN -> FAULT_ACTIVE -> RECOVER_WAIT -> RECOVER_TRY -> RUN or LOCKOUT")
    add_formula(doc, "SHORT TRIP -> LOCKOUT")

    add_section_heading(doc, "3. ADC 采样与数据换算", 1)
    add_section_heading(doc, "3.1 ADC 采样对象", 2)
    add_table(
        doc,
        ["物理量", "MCU 资源", "采样方式", "当前用途"],
        [
            ["母线电压", "ADCC4 / U3.17", "ePWM3 SOCA 触发，SOC0，采样窗 15", "参与过压判断和实时上报。"],
            ["电流（PGA 通道）", "PGA1_OF/A2 / ADCA SOC0", "ePWM3 SOCA 触发，PGA 放大 6 倍", "作为常规运行、RMS 和 I2t 主通道。"],
            ["电流（原始通道）", "B3 / ADCB SOC0", "ePWM3 SOCA 触发，不经 PGA 放大", "用于大电流和短路事件观测。"],
            ["温度", "B1 / ADCB SOC1", "ePWM3 SOCA 触发，NTC Beta 公式换算", "用于过温告警。"],
        ],
        [1.2, 1.55, 1.95, 1.8],
    )
    add_section_heading(doc, "3.2 ADC 配置与采样方式", 2)
    add_para(
        doc,
        "代码中对 ADCA、ADCB 和 ADCC 都设置了 1/4 分频，并把 ePWM3 SOCA 作为触发源，采样窗口统一为 15 个 ADC 时钟周期。PGA1 被配置为 6 倍增益。当前实现没有使用 ADC 中断和 DMA，结果是通过读取 result register 的方式获取最新值。工程按 12 位 ADC 口径做换算，但分辨率并未在当前代码里重新显式配置，说明实际取值仍应与芯片和工程初始化保持一致。",
    )
    add_table(
        doc,
        ["项目", "代码实现", "说明"],
        [
            ["触发源", "ePWM3 SOCA", "20 us 一个采样节拍。"],
            ["采样窗", "15", "对应固定的采样保持时间。"],
            ["PGA 增益", "6 倍", "用于提升正常工作电流通道分辨率。"],
            ["ADC 中断", "未发现显式实现", "当前逻辑以轮询最新结果为主。"],
            ["DMA", "未发现实现", "没有看到 ADC DMA 数据搬运。"],
            ["结果读取", "读取 RESULT 寄存器", "当前值由 sscb_adc_driver_read_latest 获取。"],
        ],
        [1.15, 3.95, 1.4],
    )
    add_section_heading(doc, "3.3 数据处理与诊断", 2)
    add_para(
        doc,
        "工程没有看到多次平均、滑动平均、一阶低通或中值滤波等通用滤波链路，当前换算主要依赖标定系数和 NTC Beta 公式。电流侧则在 protection 层使用 250 点窗口做 RMS，窗口长度对应 5 ms；如果按 20 us 采样节拍计算，250 点正好是一个完整窗口。ADC 原始值还会被简单地打上越界/饱和标志，用于自检和故障记录。",
    )
    add_table(
        doc,
        ["换算项", "当前参数 / 系数", "实现说明"],
        [
            ["母线电压", "412 mV/count", "电压原始值乘以增益后转为 0.1 V 单位。"],
            ["电流（PGA）", "224 mA/count", "常规电流主通道，输出为 0.1 A 单位。"],
            ["电流（原始）", "1343 mA/count", "大电流观测通道，输出为 0.1 A 单位。"],
            ["温度", "R_fixed=10 kOhm, R0=10 kOhm, B=3988 K", "使用 NTC Beta 公式计算摄氏温度。"],
        ],
        [1.4, 2.2, 2.9],
    )
    add_formula(doc, "V_ADC = N_ADC / (2^n - 1) * V_REF")
    add_formula(doc, "V_meas = raw_voltage * 412 mV/count / 1000")
    add_formula(doc, "I_pga = (raw_pga * 224 mA/count + offset) / 1000")
    add_formula(doc, "I_raw = (raw_raw * 1343 mA/count + offset) / 1000")
    add_formula(doc, "1/T = 1/T0 + (1/B) * ln(R_ntc / R0)")
    add_para(
        doc,
        "NTC 温度换算里，代码采用 10 kOhm、B=3988 K 的 Beta 公式，adc_raw 为 0 或接近满量程时，函数直接返回极值以避免除零。温度采样同时保留一个粗诊断标志：低于 10 或高于 4085 会被视为异常区间。",
    )

    add_section_heading(doc, "4. 保护功能设计", 1)
    add_para(
        doc,
        "保护功能已经分成短路、过载、过压和过温四类。短路依赖 CMPSS3 + ePWM Trip 硬件链路，软件只做记录、上报和闭锁；过载由 RMS + I2t 算法处理；过压按电压阈值和持续时间判断；过温当前只给告警，不直接跳闸。",
    )
    add_table(
        doc,
        ["保护功能", "默认值", "实现要点"],
        [
            ["短路", "1000 A", "CMPSS3 比较阈值换算为 DAC 码值后写入，触发后立即封锁 PWM，并进入 LOCKOUT。"],
            ["过载", "IR=250 A, tR=5000 ms, Kcool=9879", "RMS 超过整定电流后累积 I2t，达到阈值后进入 FAULT_ACTIVE。"],
            ["过压", "9600 dV / 8800 dV", "高于动作阈值持续 500 ms 后跳闸，低于恢复阈值后清零计时。"],
            ["过温", "700 dC / 650 dC", "超过 70 C 进入 WARN，降到 65 C 以下恢复。"],
        ],
        [1.0, 1.55, 3.95],
    )
    add_section_heading(doc, "4.1 短路保护", 2)
    add_para(
        doc,
        "短路保护的主链路是电流采样 -> CMPSS3 -> XBAR -> ePWM3 Trip -> PWM 封锁。软件并不把短路当成普通的延时保护处理，而是将其视为硬件快速故障。`sscb_short_threshold_to_dac_code` 会把短路整定电流换算为 CMPSS DAC 码值，当前默认参数中 `short_threshold_a = 1000`，`cmpss_dac_ref_mv = 3000`，因此软件内部的换算基准是 3.0 V。与此同时，CMPSS 驱动初始化里选择的是 VDDA 作为 DAC 参考源，因此阈值和硬件参考在联调时需要保持一致。",
    )
    add_formula(doc, "DAC_code = ShortThreshold_A * 6 * 4095 / (CmpssDacRef_mV * 10)")
    add_para(
        doc,
        "门极驱动器侧的默认配置采用较高的 DESAT 保护阈值、约 475 ns 的滤波、1.0 us 前沿消隐、故障关断方式为 TLTOff，正常开关为硬关断，并启用米勒钳位。这个配置说明驱动器级保护是后备链路，而真正的短路判据仍由 MCU 侧电流比较器完成。",
    )
    add_section_heading(doc, "4.2 过载保护", 2)
    add_para(
        doc,
        "过载保护采用 RMS + I2t 的工程化实现。代码维护一个 250 点的电流窗口，先计算整数安培级 RMS，再根据 RMS 是否超过整定电流决定是积累热量还是执行冷却衰减。I2t 触发阈值在代码里等价于 1.25 * Ir^2 * tR，因为当电流达到 1.5 Ir 时，I^2 - Ir^2 = 1.25 Ir^2。默认条件下，Ir=250 A、tR=5000 ms，因此过载保护更偏向于持续过流而非瞬态冲击。",
    )
    add_formula(doc, "I2t_acc += rms^2 - Ir^2   (when rms > Ir)")
    add_formula(doc, "I2t_acc *= 9879 / 10000   (when rms <= Ir)")
    add_formula(doc, "Trip when I2t_acc >= (2.25 * Ir^2 - Ir^2) * tR")
    add_para(
        doc,
        "故障消失后，系统不会立即恢复，而是先进入 RECOVER_WAIT，默认等待 15000 ms。只有在 Trip 清除、驱动器就绪且故障条件解除时，才会尝试恢复到 RUN；否则转入 LOCKOUT。",
    )
    add_section_heading(doc, "4.3 过压保护", 2)
    add_para(
        doc,
        "过压保护基于母线电压工程量，默认动作阈值 9600 dV，恢复阈值 8800 dV，动作延时 500 ms。当前代码里可以看到动作阈值和恢复阈值，但恢复延时参数已经定义而未在保护逻辑中直接使用；也就是说，现阶段的实现更像是“阈值回差 + 计数清零”，而不是完整的双延时回差模型。",
    )
    add_formula(doc, "V_meas >= V_trip for 500 ms -> overvoltage fault")
    add_formula(doc, "V_meas <= V_return -> clear overvoltage counter")
    add_section_heading(doc, "4.4 过温保护", 2)
    add_para(
        doc,
        "过温处理目前是告警型逻辑。NTC 采用 10 kOhm、B=3988 K 的 Beta 公式换算温度，当温度达到 70.0 C 时进入 WARN，低于 65.0 C 后清除告警。当前实现没有在过温时直接关断 PWM，这意味着它更适合作为热告警和运维提示，而不是主跳闸条件。",
    )
    add_formula(doc, "T_meas >= 70.0 C -> WARN")
    add_formula(doc, "T_meas <= 65.0 C -> clear WARN")

    add_section_heading(doc, "5. CAN 通信设计", 1)
    add_para(
        doc,
        "当前工程使用 CANB，帧格式为经典 CAN 2.0A 的 11 位标准帧，波特率 500 kbps。协议层采用设备主动周期上报和上位机命令控制的组合方式。代码把多字节字段统一按小端顺序打包，节点号默认固定为 1，参数层也限制 NodeID 只能使用 1，这意味着当前实现本质上是单节点版本。",
    )
    add_para(
        doc,
        "CAN 驱动层配置了 5 个消息对象，其中 TX 发送对象为 1，参数、时间同步和控制命令分别绑定到 2、3 和 4 号接收对象，另有 5 号故障记录对象预留。实际消息接收时，驱动会按消息对象逐个尝试读取；协议服务层目前只真正处理参数读写和手动复位命令，时间同步和故障记录读取虽然在 ID 常量里预留，但在当前协议服务中没有实现。",
    )
    add_table(
        doc,
        ["报文", "方向", "当前状态", "关键字段"],
        [
            ["0x100 + 1 实时数据", "设备 -> 上位机", "已实现", "电压、PGA 电流、温度、状态字。"],
            ["0x300 + 1 心跳", "设备 -> 上位机", "已实现", "状态码和状态字。"],
            ["0x400 + 1 参数读写", "双向", "已实现", "Op、ParamID、Value、返回状态。"],
            ["0x500 + 1 时间同步", "上位机 -> 设备", "预留", "报文 ID 已定义，当前未接入协议处理。"],
            ["0x600 + 1 控制命令", "上位机 -> 设备", "部分实现", "仅 0x01 手动复位有效。"],
            ["0x200 + 1 故障帧", "设备 -> 上位机", "接口已写，未在主循环发送", "故障码、状态和主要值。"],
            ["0x700 + 1 故障记录", "双向", "预留", "当前未见完整分包实现。"],
        ],
        [1.45, 1.05, 1.45, 2.55],
    )
    add_section_heading(doc, "5.1 实际报文字段", 2)
    add_table(
        doc,
        ["帧", "字段", "编码方式", "说明"],
        [
            ["实时帧", "Voltage_dV", "uint16, little-endian", "母线电压，0.1 V 单位。"],
            ["实时帧", "CurrentPga_dA", "int16, little-endian", "PGA 电流，0.1 A 单位。"],
            ["实时帧", "Temperature_dC", "int16, little-endian", "温度，0.1 C 单位。"],
            ["实时帧", "StatusWord", "uint16, little-endian", "系统状态字。"],
            ["心跳帧", "State", "uint8", "当前状态码。"],
            ["心跳帧", "StatusWord", "uint16, little-endian", "状态字低高字节。"],
            ["参数帧", "Op / ParamID / Value", "uint8 + uint8 + uint32", "读写操作与参数值。"],
            ["控制帧", "Cmd", "uint8", "0x01 代表手动复位。"],
        ],
        [1.15, 1.7, 1.95, 1.7],
    )

    add_section_heading(doc, "6. FRAM 参数存储设计", 1)
    add_para(
        doc,
        "硬件层已经实现了 FM25V20A 的 SPI 读写接口，地址采用 3 字节组织，写前需要 WREN，读写命令分别为 0x03 和 0x02，理论地址空间覆盖 0x00000 到 0x3FFFF。按设计文档，FRAM 用于保存参数、标定数据和故障记录，并预留启动信息和运行计数区。",
    )
    add_table(
        doc,
        ["分区", "地址范围", "容量", "用途 / 状态"],
        [
            ["BootInfo", "0x00000-0x000FF", "256 B", "魔数、版本、初始化标志，当前为设计保留。"],
            ["Parameter A", "0x00100-0x004FF", "1 KiB", "参数主副本，当前业务层未写入 FRAM。"],
            ["Parameter B", "0x00500-0x008FF", "1 KiB", "参数备份副本，当前业务层未写入 FRAM。"],
            ["Calibration", "0x00900-0x00CFF", "1 KiB", "采样标定参数，当前为设计保留。"],
            ["Runtime Counter", "0x00D00-0x00FFF", "768 B", "运行计数和掉电计数，当前为设计保留。"],
            ["Fault Ring Index", "0x01000-0x010FF", "256 B", "故障环形头，当前为设计保留。"],
            ["Fault Records", "0x01100-0x020FF", "4 KiB", "最近 50 条故障记录，当前未接入 FRAM 持久化。"],
            ["Event / Firmware Info", "0x02100 以后", "预留", "后续扩展区域。"],
        ],
        [1.25, 1.75, 0.9, 2.6],
    )
    add_para(
        doc,
        "当前代码中参数管理采用双镜像结构：sscb_param_store_t 内含 slot_a 和 slot_b，保存时在两个 RAM 镜像间轮转，启动时选择魔数、长度和校验有效且序号更大的副本。这里的完整性字段名虽然叫 crc32，但实际算法是 FNV-1a 32 位哈希，而不是标准 CRC32。若两个镜像都无效，系统会加载默认参数并回写镜像。",
    )
    add_para(
        doc,
        "需要注意的是，当前工程虽然有 FRAM 驱动和参数镜像结构，但从现有文件中没有看到参数镜像、故障日志或时间戳数据真正写入 FRAM 地址分区的调用。因此，FRAM 在当前版本更像是“接口已具备，业务持久化未完全接通”的状态。",
    )
    add_table(
        doc,
        ["结构", "字段 / 内容", "实现状态"],
        [
            ["参数镜像", "magic、version、length、sequence、crc32、sscb_params_t", "已定义，但当前保存在 RAM。"],
            ["故障日志", "sequence、fault、state、timestamp、工程量、ADC/驱动标志", "已定义，但当前保存在 RAM。"],
            ["时间基准", "unix_sec、local_ms、sync_local_ms、synced", "已定义，未见完整 CAN 同步路径。"],
            ["FRAM 低层", "SPIB 读写、WREN、3 字节地址", "已实现。"],
        ],
        [1.4, 3.15, 1.95],
    )

    add_section_heading(doc, "7. 软件运行与故障处理流程", 1)
    add_para(
        doc,
        "系统上电后先进入初始化阶段，依次完成时钟、GPIO、ADC、ePWM Trip、CMPSS、CANB、SPIB 和 I2CA 的基础准备，然后载入默认参数并尝试用参数镜像覆盖。接着执行自检：ADC 范围、NTC 粗判、CMPSS 配置、驱动器 RDYC/FLT_N 状态都会被检查；如果自检失败，系统进入 LOCKOUT。",
    )
    add_para(
        doc,
        "正常运行时，ePWM3 以 20 us 周期触发 ADC，软件读取最新结果并换算出电压、电流和温度。保护层先判断过载和过压，再判断过温告警；如果驱动器反馈短路或故障清除信号异常，则直接触发短路锁定。周期上报方面，实时数据帧按 20 ms 上报，心跳帧按 100 ms 上报，参数读写命令和手动复位命令按需处理。",
    )
    add_para(
        doc,
        "故障处理后，系统会把故障码、状态、工程量、I2t 累积值、ADC 标志和驱动器标志写入本地故障记录。过载故障在等待 15 s 后可尝试恢复；短路故障则保持锁定，必须由上位机人工复位。过温目前只产生 WARN，不强制关断；过压和过载会进入故障态并驱动状态机切换。",
    )
    add_formula(doc, "Power-on -> Init -> Self-test -> Ready -> Run")
    add_formula(doc, "Fault event -> Record -> Report -> Lockout or Recovery")

    add_section_heading(doc, "8. 总结", 1)
    add_para(
        doc,
        "当前工程已经把智能固态断路器控制器最重要的四条主线连起来了：一是 ADC 采样和工程量换算，二是以 CMPSS3 和 ePWM Trip 为核心的短路快速链路，三是以 RMS + I2t 为核心的过载保护，四是 CAN 通信与状态机管理。软件能够完成实时数据上报、参数读写和故障锁定，具备控制器软件的基本闭环特征。",
    )
    add_para(
        doc,
        "同时，工程也保留了一些后续完善点：FRAM 持久化还没有真正接入业务流程，时间同步和故障记录读取报文仍处在预留或半实现状态，NodeID 目前固定为 1，短路 DAC 参考和 CMPSS 参考源需要在联调时保持一致。就当前代码而言，这些都属于可继续迭代的工程化补强项，而不影响主体报告对现有实现的描述。",
    )

    doc.core_properties.title = "智能固态断路器控制器软件技术报告"
    doc.core_properties.subject = "智能固态断路器控制器软件技术报告"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generated from current project code and documents."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
